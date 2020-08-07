import docx

d= docx.Document ('C:\\Users\\E.H.PAUE\\python lessons\\demo.docx')
d.paragraphs

# results    :    <docx.text.paragraph.Paragraph object at 0x0000027D550AF148>, <docx.text.paragraph.Paragraph object at 0x0000027D550AFDC8>, <docx.text.paragraph.Paragraph object at 0x0000027D550AFE48>, <docx.text.paragraph.Paragraph object at 0x0000027D550AFFC8>, <docx.text.paragraph.Paragraph object at 0x0000027D550AFF88>, <docx.text.paragraph.Paragraph object at 0x0000027D550AFD48>, <docx.text.paragraph.Paragraph object at 0x0000027D550AFF48>]

d.paragraphs[0]
# result  :    <docx.text.paragraph.Paragraph object at 0x0000027D550AF408>

d.paragraphs[0].text # to get paragraph number.1

# result  :  'Document Title'  (this is the first paragraph in the word document)

d.paragraphs[1].text # to get paragraph number2

# result   : 'A plain paragraph having some bold and some italic.'   (this is the second paragraph in the word document)

p = d.paragraphs[1]
p.runs     # the run happend whenever you change the text format.(it is part of the paragraph where the test frormat change.

# results   :    [<docx.text.run.Run object at 0x00000245270AA508>, <docx.text.run.Run object at 0x00000245270AA5C8>, <docx.text.run.Run object at 0x00000245270AA788>, <docx.text.run.Run object at 0x00000245270AA7C8>]

p.runs[0].text
# result   :  'A plain paragraph having some '


p.runs[1].text
# result   :  'bold'

p.runs[2].text
# result   :  ' and some '

p.runs[3].text
# result   :   'italic.'

p.runs[3].underline= True # to activate the underline in the 4th run 

p.runs[3].text =' how are you ' # to add text  "how are you "  in the 4th run. 

d.save('C:\\Users\\E.H.PAUE\\python lessons\\demo22.docx') # to save new document with thte above changes

#________________________________________________


d= docx.Document() # it will creat emply word document insdie the python as an object .

d.add_paragraph (' hello this is a paragraph ') # to add a paragraph in this word document ' object'

d.add_paragraph (' this is paragraph number 2  ') # to add second paragraph .

d.save('C:\\Users\\E.H.PAUE\\python lessons\\demo44.docx')

p= d.paragraphs[0] # to select the first paragaraph

p.add_run ('this is a new run') # " this is new run"  will be added after the first paragraph

p.runs

p.runs[1].bold = True

#d.save()

# ________________________________________________




